"""
报表导出服务
支持PDF、Excel、Markdown、DOCX格式导出
"""

import csv
import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from sqlalchemy import and_, select
from sqlalchemy.ext.asyncio import AsyncSession

try:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.models.finance import FinancialReport, FinancialTransaction
from src.services.finance_service import FinanceService


class ReportExportService:
    """报表导出服务"""

    def __init__(self):
        pass

    async def export_to_csv(
        self,
        report_type: str,
        start_date: datetime,
        end_date: datetime,
        store_id: Optional[int] = None,
        db: Optional[AsyncSession] = None,
    ) -> bytes:
        """
        导出报表为CSV格式

        Args:
            report_type: 报表类型 (income_statement, cash_flow, balance_sheet)
            start_date: 开始日期
            end_date: 结束日期
            store_id: 门店ID
            db: 数据库会话

        Returns:
            CSV文件字节流
        """
        # 获取报表数据
        finance_service = FinanceService(db)
        if report_type == "income_statement":
            data = await finance_service.get_income_statement(start_date, end_date, store_id, db)
            return self._generate_income_statement_csv(data, start_date, end_date)
        elif report_type == "cash_flow":
            data = await finance_service.get_cash_flow_statement(start_date, end_date, store_id, db)
            return self._generate_cash_flow_csv(data, start_date, end_date)
        elif report_type == "transactions":
            data = await self._get_transactions(start_date, end_date, store_id, db)
            return self._generate_transactions_csv(data)
        else:
            raise ValueError(f"不支持的报表类型: {report_type}")

    def _generate_income_statement_csv(self, data: Dict[str, Any], start_date: datetime, end_date: datetime) -> bytes:
        """生成损益表CSV"""

        def _num(*keys: str, default: float = 0.0) -> float:
            for key in keys:
                value = data.get(key)
                if value is not None:
                    return float(value)
            return float(default)

        total_revenue = _num("total_revenue", default=_num("revenue") + _num("other_income"))
        cogs = _num("cost_of_goods_sold", "cost_of_goods")
        gross_profit = _num("gross_profit", default=total_revenue - cogs)
        total_expenses = _num(
            "total_expenses",
            "operating_expenses",
            default=_num("labor_cost") + _num("rent") + _num("utilities") + _num("marketing") + _num("other_expenses"),
        )
        operating_profit = _num("operating_profit", default=gross_profit - total_expenses)
        net_profit = _num("net_profit", default=operating_profit)
        gross_profit_margin = _num(
            "gross_profit_margin",
            default=(gross_profit / total_revenue * 100) if total_revenue else 0.0,
        )
        operating_profit_margin = _num(
            "operating_profit_margin",
            default=(operating_profit / total_revenue * 100) if total_revenue else 0.0,
        )
        net_profit_margin = _num(
            "net_profit_margin",
            default=(net_profit / total_revenue * 100) if total_revenue else 0.0,
        )

        output = io.StringIO()
        writer = csv.writer(output)

        # 写入标题
        writer.writerow(["损益表"])
        writer.writerow([f"期间: {start_date.date()} 至 {end_date.date()}"])
        writer.writerow([])

        # 写入收入部分
        writer.writerow(["收入"])
        writer.writerow(["营业收入", f"¥{_num('revenue'):,.2f}"])
        writer.writerow(["其他收入", f"¥{_num('other_income'):,.2f}"])
        writer.writerow(["总收入", f"¥{total_revenue:,.2f}"])
        writer.writerow([])

        # 写入成本部分
        writer.writerow(["成本"])
        writer.writerow(["营业成本", f"¥{cogs:,.2f}"])
        writer.writerow(["毛利润", f"¥{gross_profit:,.2f}"])
        writer.writerow(["毛利率", f"{gross_profit_margin:.2f}%"])
        writer.writerow([])

        # 写入费用部分
        writer.writerow(["费用"])
        writer.writerow(["人工成本", f"¥{_num('labor_cost'):,.2f}"])
        writer.writerow(["租金", f"¥{_num('rent'):,.2f}"])
        writer.writerow(["水电费", f"¥{_num('utilities'):,.2f}"])
        writer.writerow(["营销费用", f"¥{_num('marketing'):,.2f}"])
        writer.writerow(["其他费用", f"¥{_num('other_expenses'):,.2f}"])
        writer.writerow(["总费用", f"¥{total_expenses:,.2f}"])
        writer.writerow([])

        # 写入利润部分
        writer.writerow(["利润"])
        writer.writerow(["营业利润", f"¥{operating_profit:,.2f}"])
        writer.writerow(["营业利润率", f"{operating_profit_margin:.2f}%"])
        writer.writerow(["净利润", f"¥{net_profit:,.2f}"])
        writer.writerow(["净利润率", f"{net_profit_margin:.2f}%"])

        return output.getvalue().encode("utf-8-sig")

    def _generate_cash_flow_csv(self, data: Dict[str, Any], start_date: datetime, end_date: datetime) -> bytes:
        """生成现金流量表CSV"""

        def _num(*keys: str, default: float = 0.0) -> float:
            for key in keys:
                value = data.get(key)
                if value is not None:
                    return float(value)
            return float(default)

        output = io.StringIO()
        writer = csv.writer(output)

        # 写入标题
        writer.writerow(["现金流量表"])
        writer.writerow([f"期间: {start_date.date()} 至 {end_date.date()}"])
        writer.writerow([])

        # 经营活动现金流
        writer.writerow(["经营活动现金流"])
        writer.writerow(["销售收入", f"¥{_num('cash_from_sales'):,.2f}"])
        writer.writerow(["采购支出", f"¥{_num('cash_for_purchases'):,.2f}"])
        writer.writerow(["工资支出", f"¥{_num('cash_for_salaries'):,.2f}"])
        writer.writerow(["其他经营支出", f"¥{_num('cash_for_operations'):,.2f}"])
        writer.writerow(["经营活动净现金流", f"¥{_num('operating_cash_flow'):,.2f}"])
        writer.writerow([])

        # 投资活动现金流
        writer.writerow(["投资活动现金流"])
        writer.writerow(["设备采购", f"¥{_num('cash_for_investments'):,.2f}"])
        writer.writerow(["投资活动净现金流", f"¥{_num('investing_cash_flow'):,.2f}"])
        writer.writerow([])

        # 筹资活动现金流
        writer.writerow(["筹资活动现金流"])
        writer.writerow(["融资收入", f"¥{_num('cash_from_financing'):,.2f}"])
        writer.writerow(["筹资活动净现金流", f"¥{_num('financing_cash_flow'):,.2f}"])
        writer.writerow([])

        # 现金净变动
        writer.writerow(["现金净变动", f"¥{_num('net_cash_flow'):,.2f}"])
        writer.writerow(["期初现金", f"¥{_num('beginning_cash'):,.2f}"])
        writer.writerow(["期末现金", f"¥{_num('ending_cash'):,.2f}"])

        return output.getvalue().encode("utf-8-sig")

    def _generate_transactions_csv(self, transactions: List[Dict[str, Any]]) -> bytes:
        """生成交易明细CSV"""
        output = io.StringIO()
        writer = csv.writer(output)

        writer.writerow(["交易记录"])
        writer.writerow([])
        # 写入表头
        writer.writerow(["日期", "类型", "分类", "金额", "描述", "门店ID", "参考编号"])

        # 写入数据
        for trans in transactions:
            trans_date = trans.get("transaction_date") or trans.get("date") or datetime.utcnow()
            writer.writerow(
                [
                    trans_date.strftime("%Y-%m-%d %H:%M:%S") if hasattr(trans_date, "strftime") else str(trans_date),
                    trans.get("transaction_type", trans.get("type", "")),
                    trans.get("category", ""),
                    f"¥{float(trans.get('amount', 0)):,.2f}",
                    trans.get("description", ""),
                    trans.get("store_id", ""),
                    trans.get("reference_number", ""),
                ]
            )

        return output.getvalue().encode("utf-8-sig")

    async def _get_transactions(
        self, start_date: datetime, end_date: datetime, store_id: Optional[int], db: AsyncSession
    ) -> List[Dict[str, Any]]:
        """获取交易明细"""
        query = select(FinancialTransaction).where(
            and_(FinancialTransaction.transaction_date >= start_date, FinancialTransaction.transaction_date <= end_date)
        )

        if store_id:
            query = query.where(FinancialTransaction.store_id == store_id)

        query = query.order_by(FinancialTransaction.transaction_date.desc())

        result = await db.execute(query)
        transactions = result.scalars().all()

        return [
            {
                "transaction_date": t.transaction_date,
                "transaction_type": t.transaction_type,
                "category": t.category,
                "amount": t.amount,
                "description": t.description,
                "store_id": t.store_id,
                "reference_number": t.reference_number,
            }
            for t in transactions
        ]

    async def export_to_xlsx(
        self,
        report_type: str,
        start_date: datetime,
        end_date: datetime,
        store_id: Optional[int] = None,
        db: Optional[AsyncSession] = None,
    ) -> bytes:
        """
        导出报表为 Excel (xlsx) 格式

        Args:
            report_type: 报表类型 (income_statement, cash_flow, transactions)
            start_date: 开始日期
            end_date: 结束日期
            store_id: 门店ID
            db: 数据库会话

        Returns:
            xlsx 文件字节流
        """
        if not XLSX_AVAILABLE:
            raise ImportError("请安装 openpyxl 库以支持 Excel 导出: pip install openpyxl")

        finance_service = FinanceService(db)
        wb = openpyxl.Workbook()
        ws = wb.active

        # 通用样式
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
        section_font = Font(bold=True)
        section_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        center_align = Alignment(horizontal="center")
        right_align = Alignment(horizontal="right")

        if report_type == "income_statement":
            data = await finance_service.get_income_statement(start_date, end_date, store_id, db)
            ws.title = "损益表"
            self._write_income_statement_xlsx(
                ws, data, start_date, end_date, header_font, header_fill, section_font, section_fill, center_align, right_align
            )
        elif report_type == "cash_flow":
            data = await finance_service.get_cash_flow_statement(start_date, end_date, store_id, db)
            ws.title = "现金流量表"
            self._write_cash_flow_xlsx(
                ws, data, start_date, end_date, header_font, header_fill, section_font, section_fill, center_align, right_align
            )
        elif report_type == "transactions":
            data = await self._get_transactions(start_date, end_date, store_id, db)
            ws.title = "交易明细"
            self._write_transactions_xlsx(ws, data, header_font, header_fill, right_align)
        else:
            raise ValueError(f"不支持的报表类型: {report_type}")

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    def _write_income_statement_xlsx(
        self, ws, data, start_date, end_date, header_font, header_fill, section_font, section_fill, center_align, right_align
    ):
        """写入损益表到 Excel 工作表"""
        ws.column_dimensions["A"].width = 20
        ws.column_dimensions["B"].width = 18

        # 标题行
        ws.merge_cells("A1:B1")
        ws["A1"] = "损益表"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A1"].alignment = center_align

        ws.merge_cells("A2:B2")
        ws["A2"] = f"期间: {start_date.date()} 至 {end_date.date()}"
        ws["A2"].alignment = center_align

        rows = [
            ("收入", None, True),
            ("营业收入", data["revenue"], False),
            ("其他收入", data.get("other_income", 0), False),
            ("总收入", data["total_revenue"], False),
            (None, None, False),
            ("成本", None, True),
            ("营业成本", data["cost_of_goods_sold"], False),
            ("毛利润", data["gross_profit"], False),
            ("毛利率", f"{data['gross_profit_margin']:.2f}%", False),
            (None, None, False),
            ("费用", None, True),
            ("人工成本", data["labor_cost"], False),
            ("租金", data["rent"], False),
            ("水电费", data["utilities"], False),
            ("营销费用", data["marketing"], False),
            ("其他费用", data["other_expenses"], False),
            ("总费用", data["total_expenses"], False),
            (None, None, False),
            ("利润", None, True),
            ("营业利润", data["operating_profit"], False),
            ("营业利润率", f"{data['operating_profit_margin']:.2f}%", False),
            ("净利润", data["net_profit"], False),
            ("净利润率", f"{data['net_profit_margin']:.2f}%", False),
        ]

        for i, (label, value, is_section) in enumerate(rows, start=4):
            if label is None:
                continue
            cell_a = ws.cell(row=i, column=1, value=label)
            if is_section:
                cell_a.font = section_font
                cell_a.fill = section_fill
                ws.cell(row=i, column=2).fill = section_fill
            if value is not None:
                cell_b = ws.cell(row=i, column=2, value=f"¥{value:,.2f}" if isinstance(value, (int, float)) else value)
                cell_b.alignment = right_align

    def _write_cash_flow_xlsx(
        self, ws, data, start_date, end_date, header_font, header_fill, section_font, section_fill, center_align, right_align
    ):
        """写入现金流量表到 Excel 工作表"""
        ws.column_dimensions["A"].width = 22
        ws.column_dimensions["B"].width = 18

        ws.merge_cells("A1:B1")
        ws["A1"] = "现金流量表"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A1"].alignment = center_align

        ws.merge_cells("A2:B2")
        ws["A2"] = f"期间: {start_date.date()} 至 {end_date.date()}"
        ws["A2"].alignment = center_align

        rows = [
            ("经营活动现金流", None, True),
            ("销售收入", data["cash_from_sales"], False),
            ("采购支出", data["cash_for_purchases"], False),
            ("工资支出", data["cash_for_salaries"], False),
            ("其他经营支出", data["cash_for_operations"], False),
            ("经营活动净现金流", data["operating_cash_flow"], False),
            (None, None, False),
            ("投资活动现金流", None, True),
            ("设备采购", data["cash_for_investments"], False),
            ("投资活动净现金流", data["investing_cash_flow"], False),
            (None, None, False),
            ("筹资活动现金流", None, True),
            ("融资收入", data["cash_from_financing"], False),
            ("筹资活动净现金流", data["financing_cash_flow"], False),
            (None, None, False),
            ("现金净变动", data["net_cash_flow"], False),
            ("期初现金", data["beginning_cash"], False),
            ("期末现金", data["ending_cash"], False),
        ]

        for i, (label, value, is_section) in enumerate(rows, start=4):
            if label is None:
                continue
            cell_a = ws.cell(row=i, column=1, value=label)
            if is_section:
                cell_a.font = section_font
                cell_a.fill = section_fill
                ws.cell(row=i, column=2).fill = section_fill
            if value is not None:
                cell_b = ws.cell(row=i, column=2, value=f"¥{value:,.2f}")
                cell_b.alignment = right_align

    def _write_transactions_xlsx(self, ws, transactions, header_font, header_fill, right_align):
        """写入交易明细到 Excel 工作表"""
        headers = ["日期", "类型", "分类", "金额", "描述", "门店ID", "参考编号"]
        col_widths = [20, 12, 15, 15, 30, 12, 18]

        for col, (header, width) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            ws.column_dimensions[get_column_letter(col)].width = width

        for row, trans in enumerate(transactions, start=2):
            ws.cell(row=row, column=1, value=trans["transaction_date"].strftime("%Y-%m-%d %H:%M:%S"))
            ws.cell(row=row, column=2, value=trans["transaction_type"])
            ws.cell(row=row, column=3, value=trans["category"])
            amount_cell = ws.cell(row=row, column=4, value=trans["amount"])
            amount_cell.alignment = right_align
            ws.cell(row=row, column=5, value=trans.get("description", ""))
            ws.cell(row=row, column=6, value=str(trans.get("store_id", "")))
            ws.cell(row=row, column=7, value=trans.get("reference_number", ""))

    # ------------------------------------------------------------------ #
    # Markdown 导出                                                       #
    # ------------------------------------------------------------------ #

    async def export_to_md(
        self,
        report_type: str,
        start_date: datetime,
        end_date: datetime,
        store_id: Optional[int] = None,
        db: Optional[AsyncSession] = None,
    ) -> bytes:
        """
        导出报表为 Markdown 格式

        Returns:
            UTF-8 编码的 Markdown 字节流
        """
        finance_service = FinanceService(db)
        if report_type == "income_statement":
            data = await finance_service.get_income_statement(start_date, end_date, store_id, db)
            return self._generate_income_statement_md(data, start_date, end_date)
        elif report_type == "cash_flow":
            data = await finance_service.get_cash_flow_statement(start_date, end_date, store_id, db)
            return self._generate_cash_flow_md(data, start_date, end_date)
        elif report_type == "transactions":
            data = await self._get_transactions(start_date, end_date, store_id, db)
            return self._generate_transactions_md(data)
        else:
            raise ValueError(f"不支持的报表类型: {report_type}")

    def _generate_income_statement_md(self, data: Dict[str, Any], start_date: datetime, end_date: datetime) -> bytes:
        """生成损益表 Markdown"""

        def _num(*keys: str, default: float = 0.0) -> float:
            for key in keys:
                value = data.get(key)
                if value is not None:
                    return float(value)
            return float(default)

        total_revenue = _num("total_revenue", default=_num("revenue") + _num("other_income"))
        cogs = _num("cost_of_goods_sold", "cost_of_goods")
        gross_profit = _num("gross_profit", default=total_revenue - cogs)
        total_expenses = _num(
            "total_expenses",
            "operating_expenses",
            default=_num("labor_cost") + _num("rent") + _num("utilities") + _num("marketing") + _num("other_expenses"),
        )
        operating_profit = _num("operating_profit", default=gross_profit - total_expenses)
        net_profit = _num("net_profit", default=operating_profit)
        gross_margin = _num("gross_profit_margin", default=(gross_profit / total_revenue * 100) if total_revenue > 0 else 0.0)
        op_margin = _num("operating_profit_margin", default=(operating_profit / total_revenue * 100) if total_revenue > 0 else 0.0)
        net_margin = _num("net_profit_margin", default=(net_profit / total_revenue * 100) if total_revenue > 0 else 0.0)

        lines = [
            "# 损益表",
            "",
            f"期间：{start_date.date()} 至 {end_date.date()}",
            "",
            "## 收入",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 营业收入 | ¥{_num('revenue'):,.2f} |",
            f"| 其他收入 | ¥{_num('other_income'):,.2f} |",
            f"| **总收入** | **¥{total_revenue:,.2f}** |",
            "",
            "## 成本",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 营业成本 | ¥{cogs:,.2f} |",
            f"| **毛利润** | **¥{gross_profit:,.2f}** |",
            f"| 毛利率 | {gross_margin:.2f}% |",
            "",
            "## 费用",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 人工成本 | ¥{_num('labor_cost'):,.2f} |",
            f"| 租金 | ¥{_num('rent'):,.2f} |",
            f"| 水电费 | ¥{_num('utilities'):,.2f} |",
            f"| 营销费用 | ¥{_num('marketing'):,.2f} |",
            f"| 其他费用 | ¥{_num('other_expenses'):,.2f} |",
            f"| **总费用** | **¥{total_expenses:,.2f}** |",
            "",
            "## 利润",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 营业利润 | ¥{operating_profit:,.2f} |",
            f"| 营业利润率 | {op_margin:.2f}% |",
            f"| **净利润** | **¥{net_profit:,.2f}** |",
            f"| 净利润率 | {net_margin:.2f}% |",
        ]
        return "\n".join(lines).encode("utf-8")

    def _generate_cash_flow_md(self, data: Dict[str, Any], start_date: datetime, end_date: datetime) -> bytes:
        """生成现金流量表 Markdown"""

        def _num(*keys: str, default: float = 0.0) -> float:
            for key in keys:
                value = data.get(key)
                if value is not None:
                    return float(value)
            return float(default)

        lines = [
            "# 现金流量表",
            "",
            f"期间：{start_date.date()} 至 {end_date.date()}",
            "",
            "## 经营活动现金流",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 销售收入 | ¥{_num('cash_from_sales'):,.2f} |",
            f"| 采购支出 | ¥{_num('cash_for_purchases'):,.2f} |",
            f"| 工资支出 | ¥{_num('cash_for_salaries'):,.2f} |",
            f"| 其他经营支出 | ¥{_num('cash_for_operations'):,.2f} |",
            f"| **经营活动净现金流** | **¥{_num('operating_cash_flow'):,.2f}** |",
            "",
            "## 投资活动现金流",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 设备采购 | ¥{_num('cash_for_investments'):,.2f} |",
            f"| **投资活动净现金流** | **¥{_num('investing_cash_flow'):,.2f}** |",
            "",
            "## 筹资活动现金流",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 融资收入 | ¥{_num('cash_from_financing'):,.2f} |",
            f"| **筹资活动净现金流** | **¥{_num('financing_cash_flow'):,.2f}** |",
            "",
            "## 汇总",
            "",
            "| 项目 | 金额 |",
            "| --- | --- |",
            f"| 现金净变动 | ¥{_num('net_cash_flow'):,.2f} |",
            f"| 期初现金 | ¥{_num('beginning_cash'):,.2f} |",
            f"| **期末现金** | **¥{_num('ending_cash'):,.2f}** |",
        ]
        return "\n".join(lines).encode("utf-8")

    def _generate_transactions_md(self, transactions: List[Dict[str, Any]]) -> bytes:
        """生成交易明细 Markdown"""
        lines = [
            "# 交易记录",
            "",
            "| 日期 | 类型 | 分类 | 金额 | 描述 | 门店ID | 参考编号 |",
            "| --- | --- | --- | --- | --- | --- | --- |",
        ]
        for trans in transactions:
            trans_date = trans.get("transaction_date") or trans.get("date") or datetime.utcnow()
            date_str = trans_date.strftime("%Y-%m-%d %H:%M:%S") if hasattr(trans_date, "strftime") else str(trans_date)
            desc = str(trans.get("description", "")).replace("|", "\\|")
            lines.append(
                f"| {date_str} "
                f"| {trans.get('transaction_type', trans.get('type', ''))} "
                f"| {trans.get('category', '')} "
                f"| ¥{float(trans.get('amount', 0)):,.2f} "
                f"| {desc} "
                f"| {trans.get('store_id', '')} "
                f"| {trans.get('reference_number', '')} |"
            )
        lines.append("")
        lines.append(f"共 {len(transactions)} 条记录")
        return "\n".join(lines).encode("utf-8")

    # ------------------------------------------------------------------ #
    # DOCX 导出                                                           #
    # ------------------------------------------------------------------ #

    async def export_to_docx(
        self,
        report_type: str,
        start_date: datetime,
        end_date: datetime,
        store_id: Optional[int] = None,
        db: Optional[AsyncSession] = None,
    ) -> bytes:
        """
        导出报表为 DOCX 格式

        Returns:
            DOCX 文件字节流
        """
        if not DOCX_AVAILABLE:
            raise ImportError("请安装 python-docx 库以支持 DOCX 导出: pip install python-docx")

        finance_service = FinanceService(db)
        if report_type == "income_statement":
            data = await finance_service.get_income_statement(start_date, end_date, store_id, db)
            return self._generate_income_statement_docx(data, start_date, end_date)
        elif report_type == "cash_flow":
            data = await finance_service.get_cash_flow_statement(start_date, end_date, store_id, db)
            return self._generate_cash_flow_docx(data, start_date, end_date)
        elif report_type == "transactions":
            data = await self._get_transactions(start_date, end_date, store_id, db)
            return self._generate_transactions_docx(data)
        else:
            raise ValueError(f"不支持的报表类型: {report_type}")

    def _generate_income_statement_docx(self, data: Dict[str, Any], start_date: datetime, end_date: datetime) -> bytes:
        """生成损益表 DOCX"""

        def _num(*keys: str, default: float = 0.0) -> float:
            for key in keys:
                value = data.get(key)
                if value is not None:
                    return float(value)
            return float(default)

        total_revenue = _num("total_revenue", default=_num("revenue") + _num("other_income"))
        cogs = _num("cost_of_goods_sold", "cost_of_goods")
        gross_profit = _num("gross_profit", default=total_revenue - cogs)
        total_expenses = _num(
            "total_expenses",
            "operating_expenses",
            default=_num("labor_cost") + _num("rent") + _num("utilities") + _num("marketing") + _num("other_expenses"),
        )
        operating_profit = _num("operating_profit", default=gross_profit - total_expenses)
        net_profit = _num("net_profit", default=operating_profit)
        gross_margin = _num("gross_profit_margin", default=(gross_profit / total_revenue * 100) if total_revenue > 0 else 0.0)
        op_margin = _num("operating_profit_margin", default=(operating_profit / total_revenue * 100) if total_revenue > 0 else 0.0)
        net_margin = _num("net_profit_margin", default=(net_profit / total_revenue * 100) if total_revenue > 0 else 0.0)

        doc = DocxDocument()
        doc.add_heading("损益表", level=1)
        doc.add_paragraph(f"期间：{start_date.date()} 至 {end_date.date()}")

        # 收入部分
        doc.add_heading("收入", level=2)
        rows = [
            ("营业收入", f"¥{_num('revenue'):,.2f}"),
            ("其他收入", f"¥{_num('other_income'):,.2f}"),
            ("总收入", f"¥{total_revenue:,.2f}"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        # 成本部分
        doc.add_heading("成本", level=2)
        rows = [
            ("营业成本", f"¥{cogs:,.2f}"),
            ("毛利润", f"¥{gross_profit:,.2f}"),
            ("毛利率", f"{gross_margin:.2f}%"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        # 费用部分
        doc.add_heading("费用", level=2)
        rows = [
            ("人工成本", f"¥{_num('labor_cost'):,.2f}"),
            ("租金", f"¥{_num('rent'):,.2f}"),
            ("水电费", f"¥{_num('utilities'):,.2f}"),
            ("营销费用", f"¥{_num('marketing'):,.2f}"),
            ("其他费用", f"¥{_num('other_expenses'):,.2f}"),
            ("总费用", f"¥{total_expenses:,.2f}"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        # 利润部分
        doc.add_heading("利润", level=2)
        rows = [
            ("营业利润", f"¥{operating_profit:,.2f}"),
            ("营业利润率", f"{op_margin:.2f}%"),
            ("净利润", f"¥{net_profit:,.2f}"),
            ("净利润率", f"{net_margin:.2f}%"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def _generate_cash_flow_docx(self, data: Dict[str, Any], start_date: datetime, end_date: datetime) -> bytes:
        """生成现金流量表 DOCX"""

        def _num(*keys: str, default: float = 0.0) -> float:
            for key in keys:
                value = data.get(key)
                if value is not None:
                    return float(value)
            return float(default)

        doc = DocxDocument()
        doc.add_heading("现金流量表", level=1)
        doc.add_paragraph(f"期间：{start_date.date()} 至 {end_date.date()}")

        # 经营活动
        doc.add_heading("经营活动现金流", level=2)
        rows = [
            ("销售收入", f"¥{_num('cash_from_sales'):,.2f}"),
            ("采购支出", f"¥{_num('cash_for_purchases'):,.2f}"),
            ("工资支出", f"¥{_num('cash_for_salaries'):,.2f}"),
            ("其他经营支出", f"¥{_num('cash_for_operations'):,.2f}"),
            ("经营活动净现金流", f"¥{_num('operating_cash_flow'):,.2f}"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        # 投资活动
        doc.add_heading("投资活动现金流", level=2)
        rows = [
            ("设备采购", f"¥{_num('cash_for_investments'):,.2f}"),
            ("投资活动净现金流", f"¥{_num('investing_cash_flow'):,.2f}"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        # 筹资活动
        doc.add_heading("筹资活动现金流", level=2)
        rows = [
            ("融资收入", f"¥{_num('cash_from_financing'):,.2f}"),
            ("筹资活动净现金流", f"¥{_num('financing_cash_flow'):,.2f}"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        # 汇总
        doc.add_heading("汇总", level=2)
        rows = [
            ("现金净变动", f"¥{_num('net_cash_flow'):,.2f}"),
            ("期初现金", f"¥{_num('beginning_cash'):,.2f}"),
            ("期末现金", f"¥{_num('ending_cash'):,.2f}"),
        ]
        self._add_docx_table(doc, ["项目", "金额"], rows)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def _generate_transactions_docx(self, transactions: List[Dict[str, Any]]) -> bytes:
        """生成交易明细 DOCX"""
        doc = DocxDocument()
        doc.add_heading("交易记录", level=1)

        headers = ["日期", "类型", "分类", "金额", "描述", "门店ID", "参考编号"]
        rows = []
        for trans in transactions:
            trans_date = trans.get("transaction_date") or trans.get("date") or datetime.utcnow()
            date_str = trans_date.strftime("%Y-%m-%d %H:%M:%S") if hasattr(trans_date, "strftime") else str(trans_date)
            rows.append((
                date_str,
                trans.get("transaction_type", trans.get("type", "")),
                trans.get("category", ""),
                f"¥{float(trans.get('amount', 0)):,.2f}",
                trans.get("description", ""),
                str(trans.get("store_id", "")),
                trans.get("reference_number", ""),
            ))
        self._add_docx_table(doc, headers, rows)
        doc.add_paragraph(f"共 {len(transactions)} 条记录")

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    def _add_docx_table(self, doc, headers: List[str], rows: List[tuple]):
        """向 DOCX 文档添加带表头的表格"""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
        # 表头
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        # 数据行
        for row_idx, row_data in enumerate(rows, start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


# 全局实例
report_export_service = ReportExportService()
