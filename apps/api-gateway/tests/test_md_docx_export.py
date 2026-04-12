"""
Markdown 和 DOCX 导出功能测试
覆盖 ReportExportService 和 CustomReportService 的 MD/DOCX 导出方法

注意：本测试文件应通过正常 conftest 运行（conftest 已设置所有必要环境变量）。
如环境依赖不全，使用 --noconftest 并提供环境变量。
"""
import os
import sys

# L002: pydantic_settings 在 import 时校验环境变量，必须先设置默认值
for _k, _v in {
    "APP_ENV":               "test",
    "DATABASE_URL":          "postgresql+asyncpg://test:test@localhost/test",
    "REDIS_URL":             "redis://localhost:6379/0",
    "CELERY_BROKER_URL":     "redis://localhost:6379/0",
    "CELERY_RESULT_BACKEND": "redis://localhost:6379/0",
    "SECRET_KEY":            "test-secret-key",
    "JWT_SECRET":            "test-jwt-secret",
    "JWT_SECRET_KEY":        "test-jwt-secret-key",
}.items():
    os.environ.setdefault(_k, _v)

# 在 cryptography 不可用的环境中，mock 掉 jose/security 相关模块
# 以避免 conftest 或 import 链拉入不可用的 native 模块
from unittest.mock import MagicMock as _MagicMock
for _mod_name in [
    "jose", "jose.jwt", "jose.jws", "jose.jwk", "jose.backends",
    "jose.backends.base", "jose.backends.cryptography_backend",
]:
    sys.modules.setdefault(_mod_name, _MagicMock())

import pytest
from datetime import datetime
from unittest.mock import AsyncMock, MagicMock, patch

from src.services.report_export_service import ReportExportService
from src.services.custom_report_service import CustomReportService


# ------------------------------------------------------------------ #
# 共用测试数据                                                          #
# ------------------------------------------------------------------ #

INCOME_DATA = {
    "revenue": 100000,
    "other_income": 5000,
    "total_revenue": 105000,
    "cost_of_goods_sold": 60000,
    "gross_profit": 45000,
    "gross_profit_margin": 42.86,
    "labor_cost": 10000,
    "rent": 5000,
    "utilities": 2000,
    "marketing": 1000,
    "other_expenses": 2000,
    "total_expenses": 20000,
    "operating_profit": 25000,
    "operating_profit_margin": 23.81,
    "net_profit": 25000,
    "net_profit_margin": 23.81,
}

CASH_FLOW_DATA = {
    "cash_from_sales": 90000,
    "cash_for_purchases": 40000,
    "cash_for_salaries": 15000,
    "cash_for_operations": 5000,
    "operating_cash_flow": 30000,
    "cash_for_investments": 10000,
    "investing_cash_flow": -10000,
    "cash_from_financing": 20000,
    "financing_cash_flow": 20000,
    "net_cash_flow": 40000,
    "beginning_cash": 50000,
    "ending_cash": 90000,
}

TRANSACTIONS_DATA = [
    {
        "transaction_date": datetime(2024, 1, 15, 10, 30),
        "transaction_type": "income",
        "category": "销售",
        "amount": 1000,
        "description": "午餐销售收入",
        "store_id": "S001",
        "reference_number": "TX20240115001",
    },
    {
        "transaction_date": datetime(2024, 1, 16, 8, 0),
        "transaction_type": "expense",
        "category": "采购",
        "amount": 500,
        "description": "食材采购",
        "store_id": "S001",
        "reference_number": "TX20240116001",
    },
]

START_DATE = datetime(2024, 1, 1)
END_DATE = datetime(2024, 1, 31)


# ------------------------------------------------------------------ #
# ReportExportService — Markdown                                      #
# ------------------------------------------------------------------ #


class TestReportExportServiceMd:
    """ReportExportService Markdown 导出测试"""

    def test_generate_income_statement_md(self):
        service = ReportExportService()
        result = service._generate_income_statement_md(INCOME_DATA, START_DATE, END_DATE)

        assert isinstance(result, bytes)
        content = result.decode("utf-8")
        assert "# 损益表" in content
        assert "营业收入" in content
        assert "¥100,000.00" in content
        assert "¥25,000.00" in content
        assert "42.86%" in content
        # 确认是合法 Markdown 表格
        assert "| --- |" in content

    def test_generate_cash_flow_md(self):
        service = ReportExportService()
        result = service._generate_cash_flow_md(CASH_FLOW_DATA, START_DATE, END_DATE)

        content = result.decode("utf-8")
        assert "# 现金流量表" in content
        assert "经营活动现金流" in content
        assert "¥90,000.00" in content  # cash_from_sales
        assert "¥90,000.00" in content  # ending_cash

    def test_generate_transactions_md(self):
        service = ReportExportService()
        result = service._generate_transactions_md(TRANSACTIONS_DATA)

        content = result.decode("utf-8")
        assert "# 交易记录" in content
        assert "午餐销售收入" in content
        assert "食材采购" in content
        assert "共 2 条记录" in content

    def test_transactions_md_pipe_escape(self):
        """管道符应被转义"""
        service = ReportExportService()
        data = [
            {
                "transaction_date": datetime(2024, 1, 15),
                "transaction_type": "income",
                "category": "销售",
                "amount": 100,
                "description": "含|管道符的描述",
                "store_id": "S001",
                "reference_number": "TX001",
            }
        ]
        content = service._generate_transactions_md(data).decode("utf-8")
        assert "\\|" in content

    @pytest.mark.asyncio
    @patch("src.services.report_export_service.FinanceService")
    async def test_export_to_md_income_statement(self, mock_cls):
        service = ReportExportService()
        mock_inst = MagicMock()
        mock_inst.get_income_statement = AsyncMock(return_value=INCOME_DATA)
        mock_cls.return_value = mock_inst

        result = await service.export_to_md("income_statement", START_DATE, END_DATE)
        assert isinstance(result, bytes)
        assert "损益表" in result.decode("utf-8")

    @pytest.mark.asyncio
    @patch("src.services.report_export_service.FinanceService")
    async def test_export_to_md_cash_flow(self, mock_cls):
        service = ReportExportService()
        mock_inst = MagicMock()
        mock_inst.get_cash_flow_statement = AsyncMock(return_value=CASH_FLOW_DATA)
        mock_cls.return_value = mock_inst

        result = await service.export_to_md("cash_flow", START_DATE, END_DATE)
        assert "现金流量表" in result.decode("utf-8")

    @pytest.mark.asyncio
    async def test_export_to_md_invalid_type(self):
        service = ReportExportService()
        with pytest.raises(ValueError, match="不支持的报表类型"):
            await service.export_to_md("invalid_type", START_DATE, END_DATE)


# ------------------------------------------------------------------ #
# ReportExportService — DOCX                                          #
# ------------------------------------------------------------------ #


class TestReportExportServiceDocx:
    """ReportExportService DOCX 导出测试"""

    def test_generate_income_statement_docx(self):
        service = ReportExportService()
        result = service._generate_income_statement_docx(INCOME_DATA, START_DATE, END_DATE)

        assert isinstance(result, bytes)
        # DOCX 文件以 PK 开头（ZIP 格式）
        assert result[:2] == b"PK"

        # 解析验证内容
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "损益表" in full_text

    def test_generate_cash_flow_docx(self):
        service = ReportExportService()
        result = service._generate_cash_flow_docx(CASH_FLOW_DATA, START_DATE, END_DATE)

        assert result[:2] == b"PK"
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "现金流量表" in full_text

    def test_generate_transactions_docx(self):
        service = ReportExportService()
        result = service._generate_transactions_docx(TRANSACTIONS_DATA)

        assert result[:2] == b"PK"
        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        # 检查表格存在
        assert len(doc.tables) >= 1
        # 表头包含"日期"
        header_cells = [cell.text for cell in doc.tables[0].rows[0].cells]
        assert "日期" in header_cells

    @pytest.mark.asyncio
    @patch("src.services.report_export_service.FinanceService")
    async def test_export_to_docx_income_statement(self, mock_cls):
        service = ReportExportService()
        mock_inst = MagicMock()
        mock_inst.get_income_statement = AsyncMock(return_value=INCOME_DATA)
        mock_cls.return_value = mock_inst

        result = await service.export_to_docx("income_statement", START_DATE, END_DATE)
        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

    @pytest.mark.asyncio
    async def test_export_to_docx_invalid_type(self):
        service = ReportExportService()
        with pytest.raises(ValueError, match="不支持的报表类型"):
            await service.export_to_docx("invalid_type", START_DATE, END_DATE)


# ------------------------------------------------------------------ #
# CustomReportService — Markdown & DOCX                               #
# ------------------------------------------------------------------ #


class TestCustomReportServiceFormats:
    """CustomReportService 的 _to_md 和 _to_docx 测试"""

    COLUMNS = [
        {"field": "transaction_date", "label": "日期"},
        {"field": "category", "label": "分类"},
        {"field": "amount", "label": "金额"},
    ]
    ROWS = [
        {"transaction_date": "2024-01-15", "category": "销售", "amount": 1000},
        {"transaction_date": "2024-01-16", "category": "采购", "amount": 500},
    ]

    def test_to_md_basic(self):
        service = CustomReportService()
        result = service._to_md("测试报表", self.COLUMNS, self.ROWS)

        content = result.decode("utf-8")
        assert "# 测试报表" in content
        assert "| 日期 | 分类 | 金额 |" in content
        assert "| --- | --- | --- |" in content
        assert "1000" in content
        assert "共 2 条记录" in content

    def test_to_md_empty_rows(self):
        service = CustomReportService()
        result = service._to_md("空报表", self.COLUMNS, [])

        content = result.decode("utf-8")
        assert "暂无数据" in content

    def test_to_md_pipe_in_data(self):
        """数据中的管道符应被转义"""
        service = CustomReportService()
        rows = [{"transaction_date": "2024-01-15", "category": "A|B", "amount": 100}]
        content = service._to_md("测试", self.COLUMNS, rows).decode("utf-8")
        assert "A\\|B" in content

    def test_to_docx_basic(self):
        service = CustomReportService()
        result = service._to_docx("测试报表", self.COLUMNS, self.ROWS)

        assert isinstance(result, bytes)
        assert result[:2] == b"PK"

        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "测试报表" in full_text

        # 表格验证
        assert len(doc.tables) >= 1
        header_cells = [cell.text for cell in doc.tables[0].rows[0].cells]
        assert "日期" in header_cells
        assert "分类" in header_cells

    def test_to_docx_empty_rows(self):
        service = CustomReportService()
        result = service._to_docx("空报表", self.COLUMNS, [])

        from docx import Document
        import io

        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "暂无数据" in full_text


# ------------------------------------------------------------------ #
# ReportFormat 枚举                                                    #
# ------------------------------------------------------------------ #


class TestReportFormatEnum:
    """验证 ReportFormat 包含 MD 和 DOCX"""

    def test_md_format_exists(self):
        from src.models.report_template import ReportFormat

        assert ReportFormat.MD == "md"
        assert ReportFormat.MD.value == "md"

    def test_docx_format_exists(self):
        from src.models.report_template import ReportFormat

        assert ReportFormat.DOCX == "docx"
        assert ReportFormat.DOCX.value == "docx"
